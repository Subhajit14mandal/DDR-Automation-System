
from docx import Document
from docx.shared import Inches


def create_report(
        ddr_text,
        image_map,
        output_path):

    doc = Document()

    doc.add_heading(
        "Detailed Diagnostic Report",
        level=1
    )

    doc.add_paragraph(
        ddr_text
    )

    doc.add_page_break()

    doc.add_heading(
        "Supporting Images",
        level=1
    )

    for area, images in image_map.items():

        doc.add_heading(
            area,
            level=2
        )

        if images:

            for image in images:

                try:

                    doc.add_picture(
                        image,
                        width=Inches(3)
                    )

                except:

                    pass

        else:

            doc.add_paragraph(
                "Image Not Available"
            )

    doc.save(output_path)
